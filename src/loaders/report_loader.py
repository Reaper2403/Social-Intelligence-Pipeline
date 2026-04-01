"""
Loader: Report Generator
------------------------
Reads the AI analysis output and the master Reddit dataset, then generates
two professional Word documents (.docx) — one for post opportunities, one
for comment opportunities.

Only "Suitable" items not previously reported are included.

This is a pure Loader — it receives data from callers; it does not read
from disk itself (the caller passes pre-loaded dicts). File-system I/O for
result tracking is done via json_loader helpers.
"""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

REPORTS_DIR = Path(__file__).parent.parent.parent / "reports"


# ---------------------------------------------------------------------------
# Word document helpers (unchanged from original script)
# ---------------------------------------------------------------------------

def _add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.insert_element_before(
        pBdr,
        "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
        "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
        "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
        "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
        "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
        "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
    )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)


def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _create_report_document(report_type: str) -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading(
        f"Reddit Engagement Strategy Briefing: {report_type}", level=0
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run(
        "This document outlines high-potential engagement opportunities identified from Reddit. "
        "Each opportunity has been analyzed to provide a clear direction for response."
    ).italic = True
    score_p = doc.add_paragraph()
    score_p.add_run("Understanding the Opportunity Score: ").bold = True
    score_p.add_run(
        'This score measures the "heat" of a discussion, prioritizing newer posts and comments '
        "that are rapidly gaining upvotes and replies. A higher score indicates a more active "
        "and visible conversation."
    )
    return doc


def _populate_document(doc: Document, opportunities: list) -> None:
    first_item = True
    for opp_data in opportunities:
        if not first_item:
            _add_horizontal_line(doc.add_paragraph())
        first_item = False

        analysis = opp_data["analysis"]
        details = opp_data["details"]

        doc.add_heading(f'Opportunity Brief: {details["id"]}', level=1)
        p_info = doc.add_paragraph()
        p_info.add_run("Subreddit: ").bold = True
        p_info.add_run(f"r/{details['subreddit']} | ")
        p_info.add_run("Date: ").bold = True
        p_info.add_run(f"{details['date']}")

        p_link = doc.add_paragraph()
        p_link.add_run("Link to content: ").bold = True
        _add_hyperlink(p_link, details["url"], "Click here to view on Reddit")

        doc.add_heading("Opportunity Score", level=3)
        doc.add_paragraph(f"{details['score']:.2f}")

        doc.add_heading("Context", level=3)
        if details["type"] == "Reply to Comment":
            p_post = doc.add_paragraph()
            p_post.add_run("Original Post Title: ").bold = True
            p_post.add_run(f"{details['post_title']}\n\n")
            p_post.add_run("Original Post Body:\n").bold = True
            p_post.add_run(details["post_body"])
            p_comment = doc.add_paragraph()
            p_comment.add_run(f'\nTarget Comment by u/{details["target_author"]}:\n').bold = True
            p_comment.add_run(details["target_text"])
        else:
            doc.add_paragraph().add_run("Post Title: ").bold = True
            doc.add_paragraph(f"{details['post_title']}")
            doc.add_paragraph().add_run("Post Body:\n").bold = True
            doc.add_paragraph(details["target_text"])

        doc.add_heading("AI-Driven Strategy", level=2)
        p_strategy = doc.add_paragraph()
        p_strategy.add_run("Theme: ").bold = True
        p_strategy.add_run(f"{analysis.get('conversation_theme', 'N/A')}\n")
        p_strategy.add_run("Core Philosophy: ").bold = True
        p_strategy.add_run(f"{analysis.get('relevant_philosophy', 'N/A')}\n")
        p_strategy.add_run("Strategic Direction: ").bold = True
        p_strategy.add_run(analysis.get("strategic_direction", "N/A"))


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

class ReportLoader:
    """
    Generates Word report documents from pre-loaded data dicts.

    Args:
        full_data:            dict loaded from master_reddit_data.json
        ai_analyses:          list loaded from ai_analysis_output.json
        previously_reported:  set of opportunity IDs already written to reports
        reports_dir:          optional override for the output directory
    """

    def __init__(
        self,
        full_data: dict,
        ai_analyses: list,
        previously_reported: set,
        reports_dir: Path | None = None,
    ):
        self.full_data = full_data
        self.ai_analyses = ai_analyses
        self.previously_reported = previously_reported
        self.reports_dir = reports_dir or REPORTS_DIR
        self.reports_dir.mkdir(parents=True, exist_ok=True)

        # Build fast lookup tables
        self._post_lookup = {
            f"post_{p['post_details']['id']}": p
            for p in full_data.get("posts", [])
            if "id" in p.get("post_details", {})
        }
        self._comment_lookup = {
            f"comment_{c['id']}": (c, p)
            for p in full_data.get("posts", [])
            for c in p.get("top_comments", [])
            if "id" in c
        }

    def generate(self) -> set:
        """
        Filter for suitable, unreported opportunities and emit Word documents.

        Returns:
            set of newly reported opportunity IDs (so the caller can persist them)
        """
        post_opportunities: list = []
        comment_opportunities: list = []
        newly_reported: set = set()
        skipped_unsuitable = 0
        skipped_prev_reported = 0

        for analysis in self.ai_analyses:
            opp_id = analysis.get("opportunity_id")
            if not opp_id:
                continue
            if analysis.get("status") != "Suitable":
                skipped_unsuitable += 1
                continue
            if opp_id in self.previously_reported:
                skipped_prev_reported += 1
                continue

            if opp_id.startswith("post_") and opp_id in self._post_lookup:
                post = self._post_lookup[opp_id]["post_details"]
                details = {
                    "id": opp_id,
                    "type": "Reply to Post",
                    "url": post.get("url"),
                    "subreddit": post.get("subreddit"),
                    "post_title": post.get("title"),
                    "target_author": post.get("author"),
                    "score": post.get("opportunity_score_post", 0),
                    "target_text": post.get("body"),
                    "post_body": post.get("body", ""),
                    "date": post.get("created_utc", " ").split(" ")[0],
                }
                post_opportunities.append({"analysis": analysis, "details": details})
                newly_reported.add(opp_id)

            elif opp_id.startswith("comment_") and opp_id in self._comment_lookup:
                comment, post_data = self._comment_lookup[opp_id]
                post_details = post_data["post_details"]
                details = {
                    "id": opp_id,
                    "type": "Reply to Comment",
                    "url": post_details.get("url"),
                    "subreddit": post_details.get("subreddit"),
                    "post_title": post_details.get("title"),
                    "target_author": comment.get("author"),
                    "score": comment.get("opportunity_score_reply", 0),
                    "target_text": comment.get("body"),
                    "post_body": post_details.get("body", ""),
                    "date": comment.get("created_utc", " ").split(" ")[0],
                }
                comment_opportunities.append({"analysis": analysis, "details": details})
                newly_reported.add(opp_id)

        logger.info(
            f"Filtering done: {skipped_unsuitable} Unsuitable, "
            f"{skipped_prev_reported} previously reported."
        )

        if post_opportunities:
            doc = _create_report_document("Post Opportunities")
            _populate_document(doc, post_opportunities)
            out = self.reports_dir / "Report_Posts.docx"
            doc.save(out)
            logger.info(f"Post report saved → {out} ({len(post_opportunities)} items)")
        else:
            logger.info("No new suitable post opportunities.")

        if comment_opportunities:
            doc = _create_report_document("Comment Opportunities")
            _populate_document(doc, comment_opportunities)
            out = self.reports_dir / "Report_Comments.docx"
            doc.save(out)
            logger.info(f"Comment report saved → {out} ({len(comment_opportunities)} items)")
        else:
            logger.info("No new suitable comment opportunities.")

        return newly_reported
