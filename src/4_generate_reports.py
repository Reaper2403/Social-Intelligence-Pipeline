"""
Stage 4: Generate Reports
--------------------------
Reads the AI analysis output and the master Reddit dataset, then generates two professional
Word documents (.docx) — one for post opportunities, one for comment opportunities.
Only "Suitable" items not previously reported are included. Tracks reported IDs to prevent
duplicate reports on subsequent runs.

Inputs:  data/master_reddit_data.json, data/ai_analysis_output.json, data/reported_ids.log
Outputs: reports/Report_Posts.docx, reports/Report_Comments.docx, data/reported_ids.log (updated)

Usage:
    python src/4_generate_reports.py
"""

import json
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# --- Directory paths relative to this script's location ---
DATA_DIR = Path(__file__).parent.parent / "data"
REPORTS_DIR = Path(__file__).parent.parent / "reports"

FULL_DATA_FILE = DATA_DIR / "master_reddit_data.json"
AI_OUTPUT_FILE = DATA_DIR / "ai_analysis_output.json"
HISTORY_LOG_FILE = DATA_DIR / "reported_ids.log"
POST_REPORT_FILE = REPORTS_DIR / "Report_Posts.docx"
COMMENT_REPORT_FILE = REPORTS_DIR / "Report_Comments.docx"


def add_horizontal_line(paragraph):
    """Adds a horizontal line to a paragraph for separation."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_hyperlink(paragraph, url, text):
    """
    Appends a hyperlink to a paragraph object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style the link
    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0563C1") # Standard blue hyperlink color
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def create_report_document(report_type):
    """Creates and styles a new Word document for the report."""
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    doc.add_heading(f'Reddit Engagement Strategy Briefing: {report_type}', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run('This document outlines high-potential engagement opportunities identified from Reddit. Each opportunity has been analyzed to provide a clear direction for response.').italic = True

    score_p = doc.add_paragraph()
    score_p.add_run('Understanding the Opportunity Score: ').bold = True
    score_p.add_run('This score measures the "heat" of a discussion, prioritizing newer posts and comments that are rapidly gaining upvotes and replies. A higher score indicates a more active and visible conversation.')

    return doc

def populate_document(doc, opportunities):
    """Populates a document with a list of opportunities, including clickable links."""
    first_item = True
    for opp_data in opportunities:
        if not first_item:
            add_horizontal_line(doc.add_paragraph())
        first_item = False

        analysis = opp_data['analysis']
        details = opp_data['details']

        doc.add_heading(f'Opportunity Brief: {details["id"]}', level=1)

        p_info = doc.add_paragraph()
        p_info.add_run(f"Subreddit: ").bold = True
        p_info.add_run(f"r/{details['subreddit']} | ")
        p_info.add_run(f"Date: ").bold = True
        p_info.add_run(f"{details['date']}")

        p_link = doc.add_paragraph()
        p_link.add_run("Link to content: ").bold = True
        add_hyperlink(p_link, details['url'], "Click here to view on Reddit")

        doc.add_heading('Opportunity Score', level=3)
        doc.add_paragraph(f"{details['score']:.2f}")

        doc.add_heading('Context', level=3)
        if details['type'] == "Reply to Comment":
            p_post = doc.add_paragraph()
            p_post.add_run('Original Post Title: ').bold = True
            p_post.add_run(f"{details['post_title']}\n\n")
            p_post.add_run('Original Post Body:\n').bold = True
            p_post.add_run(details['post_body'])

            p_comment = doc.add_paragraph()
            p_comment.add_run(f'\nTarget Comment by u/{details["target_author"]}:\n').bold = True
            p_comment.add_run(details['target_text'])
        else: # Reply to Post
            doc.add_paragraph().add_run('Post Title: ').bold = True
            doc.add_paragraph(f"{details['post_title']}")
            doc.add_paragraph().add_run('Post Body:\n').bold = True
            doc.add_paragraph(details['target_text'])

        doc.add_heading('AI-Driven Strategy', level=2)
        p_strategy = doc.add_paragraph()
        p_strategy.add_run('Theme: ').bold = True
        p_strategy.add_run(f"{analysis.get('conversation_theme', 'N/A')}\n")
        p_strategy.add_run('Core Philosophy: ').bold = True
        p_strategy.add_run(f"{analysis.get('relevant_philosophy', 'N/A')}\n")
        p_strategy.add_run('Strategic Direction: ').bold = True
        p_strategy.add_run(analysis.get('strategic_direction', 'N/A'))

def generate_final_reports_with_links():
    """
    Filters for suitable, new opportunities and generates two professional
    Word documents with clickable links, while tracking reported IDs.
    """
    logger.info("Generating Final, Filtered Word Reports with Clickable Links")

    REPORTS_DIR.mkdir(exist_ok=True)

    try:
        with open(HISTORY_LOG_FILE, 'r', encoding='utf-8') as f:
            previously_reported_ids = set(line.strip() for line in f)
        logger.info(f"Loaded {len(previously_reported_ids)} IDs from the report history log.")
    except FileNotFoundError:
        previously_reported_ids = set()
        logger.info("No report history log found. A new one will be created.")

    try:
        with open(FULL_DATA_FILE, 'r', encoding='utf-8') as f:
            full_data = json.load(f)
        with open(AI_OUTPUT_FILE, 'r', encoding='utf-8') as f:
            ai_analyses = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError) as e:
        logger.error(f"Could not load required files. Error: {e}")
        return

    post_lookup = {f"post_{p['post_details']['id']}": p for p in full_data.get("posts", []) if 'id' in p.get('post_details', {})}
    comment_lookup = {f"comment_{c['id']}": (c, p) for p in full_data.get("posts", []) for c in p.get("top_comments", []) if 'id' in c}

    post_opportunities = []
    comment_opportunities = []
    newly_reported_ids = set()
    skipped_unsuitable = 0
    skipped_previously_reported = 0

    for analysis in ai_analyses:
        opp_id = analysis.get("opportunity_id")
        if not opp_id:
            continue

        if analysis.get("status") != "Suitable":
            skipped_unsuitable += 1
            continue

        if opp_id in previously_reported_ids:
            skipped_previously_reported += 1
            continue

        details = {}
        if opp_id.startswith("post_") and opp_id in post_lookup:
            post = post_lookup[opp_id]['post_details']
            details = { "id": opp_id, "type": "Reply to Post", "url": post.get('url'), "subreddit": post.get('subreddit'), "post_title": post.get('title'), "target_author": post.get('author'), "score": post.get('opportunity_score_post', 0), "target_text": post.get('body'), "date": post.get('created_utc', ' ').split(' ')[0] }
            post_opportunities.append({"analysis": analysis, "details": details})
            newly_reported_ids.add(opp_id)

        elif opp_id.startswith("comment_") and opp_id in comment_lookup:
            comment, post_data = comment_lookup[opp_id]
            post_details = post_data['post_details']
            details = { "id": opp_id, "type": "Reply to Comment", "url": post_details.get('url'), "subreddit": post_details.get('subreddit'), "post_title": post_details.get('title'), "target_author": comment.get('author'), "score": comment.get('opportunity_score_reply', 0), "target_text": comment.get('body'), "post_body": post_details.get('body', ''), "date": comment.get('created_utc', ' ').split(' ')[0] }
            comment_opportunities.append({"analysis": analysis, "details": details})
            newly_reported_ids.add(opp_id)

    logger.info(f"Filtering complete: {skipped_unsuitable} Unsuitable skipped, {skipped_previously_reported} previously reported skipped.")

    if post_opportunities:
        post_doc = create_report_document("Post Opportunities")
        populate_document(post_doc, post_opportunities)
        post_doc.save(POST_REPORT_FILE)
        logger.info(f"Generated '{POST_REPORT_FILE}' with {len(post_opportunities)} new post opportunities.")
    else:
        logger.info("No new suitable post opportunities found to generate a report.")

    if comment_opportunities:
        comment_doc = create_report_document("Comment Opportunities")
        populate_document(comment_doc, comment_opportunities)
        comment_doc.save(COMMENT_REPORT_FILE)
        logger.info(f"Generated '{COMMENT_REPORT_FILE}' with {len(comment_opportunities)} new comment opportunities.")
    else:
        logger.info("No new suitable comment opportunities found to generate a report.")

    if newly_reported_ids:
        with open(HISTORY_LOG_FILE, 'a', encoding='utf-8') as f:
            for item_id in sorted(list(newly_reported_ids)):
                f.write(f"{item_id}\n")
        logger.info(f"Updated '{HISTORY_LOG_FILE}' with {len(newly_reported_ids)} new IDs.")

if __name__ == "__main__":
    generate_final_reports_with_links()
